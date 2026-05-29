"""
DocumentParser: reads .docx files and returns a structured ParsedDocument
with sections, headings, and full text.
"""
from __future__ import annotations
from pathlib import Path
from typing import List, Optional
from pydantic import BaseModel, Field


class DocumentSection(BaseModel):
    heading: str
    level: int          # 1, 2, or 3 (Word heading level)
    content: str
    page_hint: Optional[int] = None


class ParsedDocument(BaseModel):
    filename: str
    full_text: str
    sections: List[DocumentSection] = Field(default_factory=list)
    word_count: int = 0


class DocumentParser:
    """Parses .docx files into structured sections."""

    def parse(self, file_path: str) -> ParsedDocument:
        from docx import Document

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        doc = Document(str(path))
        sections: List[DocumentSection] = []
        current_heading = "Introduction"
        current_level = 1
        current_content: List[str] = []
        full_text_parts: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            full_text_parts.append(text)

            # Detect heading styles
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                # Save previous section
                if current_content:
                    sections.append(DocumentSection(
                        heading=current_heading,
                        level=current_level,
                        content="\n".join(current_content),
                    ))
                    current_content = []

                # Determine heading level
                try:
                    level = int(style_name.split()[-1])
                except ValueError:
                    level = 1

                current_heading = text
                current_level = level
            else:
                current_content.append(text)

        # Flush last section
        if current_content:
            sections.append(DocumentSection(
                heading=current_heading,
                level=current_level,
                content="\n".join(current_content),
            ))

        full_text = "\n".join(full_text_parts)
        return ParsedDocument(
            filename=path.name,
            full_text=full_text,
            sections=sections,
            word_count=len(full_text.split()),
        )

    def get_section_context(self, document: ParsedDocument, query_terms: List[str]) -> str:
        """
        Returns concatenated content from sections whose headings or content
        contain any of the query_terms (case-insensitive).
        """
        query_lower = [t.lower() for t in query_terms]
        matched_parts: List[str] = []

        for section in document.sections:
            heading_lower = section.heading.lower()
            content_lower = section.content.lower()
            if any(
                term in heading_lower or term in content_lower
                for term in query_lower
            ):
                matched_parts.append(f"[{section.heading}]\n{section.content}")

        return "\n\n".join(matched_parts) if matched_parts else ""
