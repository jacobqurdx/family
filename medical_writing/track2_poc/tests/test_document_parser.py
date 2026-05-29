"""
Tests for ingestion/document_parser.py
"""
from __future__ import annotations
import sys
import os
from pathlib import Path

# Ensure project root is on sys.path
sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest
import tempfile
from docx import Document

from ingestion.document_parser import DocumentParser, ParsedDocument


def make_mini_docx(tmp_path: Path) -> Path:
    """Create a minimal .docx with known structure."""
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the introduction paragraph about type 2 diabetes mellitus.")
    doc.add_heading("Study Design", level=2)
    doc.add_paragraph("A randomized double-blind placebo-controlled study.")
    doc.add_paragraph("The investigational drug is STR-4021.")
    doc.add_heading("Objectives", level=2)
    doc.add_paragraph("To evaluate the efficacy of STR-4021.")
    path = tmp_path / "test_protocol.docx"
    doc.save(str(path))
    return path


def test_parse_mini_docx(tmp_path):
    parser = DocumentParser()
    doc_path = make_mini_docx(tmp_path)
    parsed = parser.parse(str(doc_path))

    assert isinstance(parsed, ParsedDocument)
    assert parsed.filename == "test_protocol.docx"
    assert parsed.word_count > 0
    assert len(parsed.sections) > 0
    assert parsed.full_text != ""


def test_section_headings(tmp_path):
    parser = DocumentParser()
    doc_path = make_mini_docx(tmp_path)
    parsed = parser.parse(str(doc_path))

    headings = [s.heading for s in parsed.sections]
    assert "Introduction" in headings
    assert "Study Design" in headings


def test_section_content(tmp_path):
    parser = DocumentParser()
    doc_path = make_mini_docx(tmp_path)
    parsed = parser.parse(str(doc_path))

    design_section = next((s for s in parsed.sections if s.heading == "Study Design"), None)
    assert design_section is not None
    assert "STR-4021" in design_section.content or "randomized" in design_section.content


def test_get_section_context_found(tmp_path):
    parser = DocumentParser()
    doc_path = make_mini_docx(tmp_path)
    parsed = parser.parse(str(doc_path))

    context = parser.get_section_context(parsed, ["STR-4021"])
    assert "STR-4021" in context


def test_get_section_context_not_found(tmp_path):
    parser = DocumentParser()
    doc_path = make_mini_docx(tmp_path)
    parsed = parser.parse(str(doc_path))

    context = parser.get_section_context(parsed, ["TOTALLY_NONEXISTENT_TERM_XYZ"])
    assert context == ""


def test_get_section_context_by_heading(tmp_path):
    parser = DocumentParser()
    doc_path = make_mini_docx(tmp_path)
    parsed = parser.parse(str(doc_path))

    context = parser.get_section_context(parsed, ["objectives"])
    assert "efficacy" in context.lower() or "STR-4021" in context


def test_parse_synth_protocol():
    """Parse the generated synthetic protocol if it exists."""
    proto_path = Path(__file__).parent.parent / "data" / "protocols" / "synth_protocol.docx"
    if not proto_path.exists():
        pytest.skip("synth_protocol.docx not yet generated — run scripts/generate_synth_protocol.py")

    parser = DocumentParser()
    parsed = parser.parse(str(proto_path))

    assert parsed.word_count > 1000
    assert len(parsed.sections) >= 5
    assert "STR-4021" in parsed.full_text
    assert "diabetes" in parsed.full_text.lower()


def test_parse_synth_protocol_sections():
    """Verify key sections are present in synthetic protocol."""
    proto_path = Path(__file__).parent.parent / "data" / "protocols" / "synth_protocol.docx"
    if not proto_path.exists():
        pytest.skip("synth_protocol.docx not yet generated.")

    parser = DocumentParser()
    parsed = parser.parse(str(proto_path))

    full_lower = parsed.full_text.lower()
    assert "hba1c" in full_lower
    assert "metformin" in full_lower
    assert "randomized" in full_lower or "randomised" in full_lower
