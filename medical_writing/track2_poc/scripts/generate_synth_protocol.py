"""
Generate data/protocols/synth_protocol.docx — a realistic synthetic clinical
trial protocol for STR-4021-201 that covers all 57 USDM graph nodes.

Run from the track2_poc directory:
    python scripts/generate_synth_protocol.py
"""
from __future__ import annotations
import sys
from pathlib import Path

# Allow running from track2_poc directory
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


def add_bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def generate_protocol():
    doc = Document()

    # ─────────────────────────────────────────────────
    # TITLE PAGE
    # ─────────────────────────────────────────────────
    title = doc.add_heading("", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "A Phase 2, Randomized, Double-Blind, Placebo-Controlled, "
        "Parallel-Group Study to Evaluate the Efficacy and Safety of STR-4021 "
        "in Adults with Type 2 Diabetes Mellitus with Inadequate Glycemic "
        "Control on Metformin Monotherapy"
    )
    run.bold = True

    doc.add_paragraph("")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Sponsor: Structure Therapeutics\n").bold = True
    meta.add_run("Protocol Number: STR-4021-201\n")
    meta.add_run("Version 1.0\n")
    meta.add_run("Date: 15 January 2026\n")
    meta.add_run("EudraCT Number: 2026-000123-42\n")
    meta.add_run("IND Number: 123456\n")

    doc.add_page_break()

    # ─────────────────────────────────────────────────
    # SECTION 1 — INTRODUCTION AND BACKGROUND
    # ─────────────────────────────────────────────────
    add_heading(doc, "1. Introduction and Background", level=1)
    add_paragraph(doc,
        "Type 2 diabetes mellitus (T2DM) is a chronic metabolic disorder characterised by "
        "progressive beta-cell dysfunction and insulin resistance. Despite the availability of "
        "numerous antidiabetic agents, a substantial proportion of patients with T2DM fail to "
        "achieve glycaemic targets on metformin monotherapy, creating a significant unmet need "
        "for additional treatment options."
    )
    add_paragraph(doc,
        "STR-4021 is a novel, orally administered, once-daily investigational compound developed "
        "by Structure Therapeutics for the treatment of type 2 diabetes mellitus with inadequate "
        "glycemic control on metformin monotherapy. Preclinical studies have demonstrated that "
        "STR-4021 acts as a potent and selective agonist of the GLP-1 receptor, promoting insulin "
        "secretion in a glucose-dependent manner."
    )

    add_heading(doc, "1.1 Indication and Unmet Need", level=2)
    add_paragraph(doc,
        "The indication under investigation is type 2 diabetes mellitus with inadequate glycemic "
        "control on metformin monotherapy. HbA1c remains the primary surrogate biomarker for "
        "long-term glycaemic control and is recognised by regulatory authorities including the FDA "
        "and EMA as the primary endpoint for T2DM drug approval. The primary endpoint for this "
        "study is change from baseline in HbA1c at Week 24."
    )

    # ─────────────────────────────────────────────────
    # SECTION 2 — STUDY OBJECTIVES
    # ─────────────────────────────────────────────────
    add_heading(doc, "2. Study Objectives and Endpoints", level=1)

    add_heading(doc, "2.1 Primary Objective", level=2)
    add_paragraph(doc,
        "To evaluate the efficacy of STR-4021 compared to placebo on glycemic control as measured "
        "by change from baseline in HbA1c at Week 24 in adults with type 2 diabetes mellitus "
        "inadequately controlled on metformin monotherapy."
    )

    add_heading(doc, "2.2 Secondary Objectives", level=2)
    add_bullet(doc, "To evaluate the effect of STR-4021 on fasting plasma glucose (FPG) at Week 24.")
    add_bullet(doc, "To assess the proportion of patients achieving HbA1c <7.0% at Week 24.")
    add_bullet(doc, "To evaluate body weight change from baseline at Week 24.")
    add_bullet(doc, "To assess the safety and tolerability of STR-4021.")

    add_heading(doc, "2.3 Exploratory Objectives", level=2)
    add_bullet(doc, "To evaluate biomarkers of insulin secretion including C-peptide and proinsulin.")
    add_bullet(doc, "To explore patient-reported outcomes using the Diabetes Treatment Satisfaction Questionnaire (DTSQ).")

    add_heading(doc, "2.4 Primary Endpoint", level=2)
    add_paragraph(doc,
        "Primary Endpoint: Change from baseline in HbA1c (%) at Week 24, measured by central "
        "laboratory HbA1c assay (NGSP-certified). The primary endpoint timepoint is Week 24."
    )

    add_heading(doc, "2.5 Secondary Endpoints", level=2)
    add_bullet(doc, "Change from baseline in fasting plasma glucose (mmol/L) at Week 24.")
    add_bullet(doc, "Proportion of patients achieving HbA1c <7.0% at Week 24.")
    add_bullet(doc, "Change from baseline in body weight (kg) at Week 24.")
    add_bullet(doc, "Change from baseline in 2-hour postprandial glucose at Week 24.")

    # ─────────────────────────────────────────────────
    # SECTION 3 — STUDY DESIGN
    # ─────────────────────────────────────────────────
    add_heading(doc, "3. Study Design", level=1)

    add_heading(doc, "3.1 Overall Design", level=2)
    add_paragraph(doc,
        "This is a Phase 2, randomized, double-blind, placebo-controlled, parallel-group study. "
        "The intervention model is parallel group with three treatment arms. The blinding type is "
        "double-blind, with subjects, investigators, and sponsor personnel blinded to treatment "
        "assignment. The randomization type is stratified block randomization."
    )
    add_paragraph(doc,
        "Eligible subjects will be randomized in a 1:1:1 ratio to one of three treatment arms:"
    )
    add_bullet(doc, "STR-4021 10 mg once daily")
    add_bullet(doc, "STR-4021 25 mg once daily")
    add_bullet(doc, "Placebo once daily (matching placebo tablet)")

    add_heading(doc, "3.2 Study Epochs", level=2)
    add_paragraph(doc,
        "The study comprises three sequential epochs:"
    )
    add_bullet(doc, "Screening (4 weeks): Assessment of eligibility criteria.")
    add_bullet(doc, "Treatment (24 weeks): Double-blind administration of study drug.")
    add_bullet(doc, "Follow-up (4 weeks): Safety follow-up after treatment completion.")
    add_paragraph(doc,
        "The total study duration is 28 weeks from first visit to last visit. "
        "The treatment duration is 24 weeks. The follow-up duration is 4 weeks."
    )

    add_heading(doc, "3.3 Randomization and Stratification", level=2)
    add_paragraph(doc,
        "Subjects will be randomized using a central IRT system with permuted block randomization "
        "in a 1:1:1 ratio. Stratification factors are: (1) HbA1c at screening (<8.5% vs ≥8.5%) "
        "and (2) Geographic region (North America vs Europe vs Asia-Pacific)."
    )

    add_heading(doc, "3.4 Study Elements", level=2)
    add_paragraph(doc, "The following study elements are included within the epochs:")
    for element in [
        "Screening Visit (Week -4)",
        "Randomization/Baseline (Week 0)",
        "Week 4 Visit",
        "Week 8 Visit",
        "Week 12 Visit",
        "Week 16 Visit",
        "Week 20 Visit",
        "End of Treatment (Week 24)",
        "Follow-up Visit (Week 28)",
    ]:
        add_bullet(doc, element)

    # ─────────────────────────────────────────────────
    # SECTION 4 — STUDY POPULATION
    # ─────────────────────────────────────────────────
    add_heading(doc, "4. Study Population", level=1)

    add_heading(doc, "4.1 Population Description", level=2)
    add_paragraph(doc,
        "The study population consists of adult patients aged 18 to 75 years with type 2 diabetes "
        "mellitus inadequately controlled on stable metformin monotherapy. Both male and female "
        "subjects are eligible. Approximately 240 subjects will be enrolled at approximately 45 "
        "sites globally, with 80 subjects per arm."
    )

    add_heading(doc, "4.2 Inclusion Criteria", level=2)
    add_paragraph(doc, "Subjects must meet all of the following criteria:")
    for criterion in [
        "Age 18-75 years inclusive",
        "Diagnosis of type 2 diabetes mellitus for at least 6 months prior to screening",
        "HbA1c ≥7.5% and ≤10.5% at screening",
        "On stable metformin monotherapy (≥1000 mg/day) for at least 3 months prior to screening",
        "BMI 22-42 kg/m² at screening",
    ]:
        add_bullet(doc, criterion)

    add_heading(doc, "4.3 Exclusion Criteria", level=2)
    add_paragraph(doc, "Subjects must not meet any of the following criteria:")
    for criterion in [
        "Type 1 diabetes mellitus",
        "Use of any antidiabetic agent other than metformin within 3 months of screening",
        "eGFR <45 mL/min/1.73m² at screening",
        "History of severe hypoglycemia within 6 months of screening",
        "Active cardiovascular disease within 6 months of screening",
    ]:
        add_bullet(doc, criterion)

    # ─────────────────────────────────────────────────
    # SECTION 5 — TREATMENT
    # ─────────────────────────────────────────────────
    add_heading(doc, "5. Treatment Administration", level=1)

    add_heading(doc, "5.1 Investigational Product", level=2)
    add_paragraph(doc,
        "The investigational product is STR-4021, available as 10 mg and 25 mg oral tablets. "
        "The dose unit is mg. The route of administration is oral. Subjects will receive "
        "once daily dosing. Study drug should be administered orally once daily in the morning "
        "with or without food."
    )

    add_heading(doc, "5.2 Comparator", level=2)
    add_paragraph(doc,
        "The comparator is matching placebo tablet administered orally once daily in the morning. "
        "The placebo tablet is identical in appearance to the STR-4021 tablets."
    )

    add_heading(doc, "5.3 Dose Modifications", level=2)
    add_paragraph(doc, "The following dose modifications are permitted:")
    add_bullet(doc, "Dose interruption for serious adverse events (SAEs) until resolved to baseline or acceptable level.")
    add_bullet(doc, "Permanent discontinuation for severe hypersensitivity reactions.")

    # ─────────────────────────────────────────────────
    # SECTION 6 — SCHEDULE OF ACTIVITIES
    # ─────────────────────────────────────────────────
    add_heading(doc, "6. Schedule of Activities", level=1)

    add_heading(doc, "6.1 Visit Schedule", level=2)
    add_paragraph(doc, "The following visits are included in the Schedule of Activities:")
    for visit in [
        "Screening (Week -4)",
        "Baseline/Randomization (Week 0)",
        "Week 4",
        "Week 8",
        "Week 12",
        "Week 16",
        "Week 20",
        "End of Treatment (Week 24)",
        "Follow-up (Week 28)",
    ]:
        add_bullet(doc, visit)

    add_heading(doc, "6.2 Activities", level=2)
    add_paragraph(doc, "The following assessments and activities will be performed at scheduled visits:")
    for activity in [
        "Informed Consent",
        "Eligibility Assessment",
        "Medical History",
        "Physical Examination",
        "Vital Signs (blood pressure, heart rate, temperature)",
        "12-lead ECG",
        "Central Lab (HbA1c, FPG, lipids, eGFR, chemistry panel)",
        "OGTT/MMTT (selected visits)",
        "Study Drug Dispensing and Accountability",
        "Concomitant Medications Review",
        "Adverse Event Assessment",
        "Body Weight",
        "Waist Circumference",
    ]:
        add_bullet(doc, activity)

    # ─────────────────────────────────────────────────
    # SECTION 7 — STATISTICAL METHODS
    # ─────────────────────────────────────────────────
    add_heading(doc, "7. Statistical Methods", level=1)

    add_heading(doc, "7.1 Analysis Populations", level=2)
    add_paragraph(doc, "The following analysis populations are defined:")
    add_bullet(doc,
        "Full Analysis Set (FAS): all randomized subjects who received at least one dose of "
        "study drug and had at least one post-baseline HbA1c measurement. This is the primary "
        "analysis population."
    )
    add_bullet(doc,
        "Per Protocol Set (PPS): FAS subjects without major protocol deviations likely to "
        "influence the primary endpoint assessment."
    )
    add_bullet(doc,
        "Safety Population: all randomized subjects who received at least one dose of study drug."
    )

    add_heading(doc, "7.2 Primary Analysis", level=2)
    add_paragraph(doc,
        "The primary analysis will use a Mixed Model for Repeated Measures (MMRM) with change "
        "from baseline in HbA1c as the dependent variable. The MMRM model will include fixed "
        "effects for treatment arm, visit, treatment-by-visit interaction, baseline HbA1c, and "
        "stratification factors (baseline HbA1c category and geographic region). An unstructured "
        "covariance matrix will be used. REML estimation will be applied."
    )
    add_paragraph(doc,
        "The primary analysis type is Mixed Model for Repeated Measures (MMRM). The full "
        "description of the primary statistical analysis method: The MMRM will be fitted using "
        "SAS PROC MIXED with the REPEATED statement and TYPE=UN covariance structure. "
        "Kenward-Roger degrees of freedom will be used for all inferential testing."
    )

    add_heading(doc, "7.3 Missing Data", level=2)
    add_paragraph(doc,
        "Missing data are handled implicitly by the MMRM model under the missing at random (MAR) "
        "assumption. Pre-specified sensitivity analyses using pattern mixture models will be "
        "conducted to assess robustness under missing not at random (MNAR) assumptions."
    )

    add_heading(doc, "7.4 Multiplicity", level=2)
    add_paragraph(doc,
        "A hierarchical testing procedure will be used for secondary endpoints in order of "
        "pre-specified priority to control the familywise type I error rate at 0.05 (two-sided)."
    )

    add_heading(doc, "7.5 Estimands", level=2)
    add_paragraph(doc,
        "The treatment policy estimand (per ICH E9(R1)) is the primary estimand, targeting the "
        "effect of treatment assignment regardless of treatment discontinuation or rescue medication use."
    )

    add_heading(doc, "7.6 Sample Size", level=2)
    add_paragraph(doc,
        "A sample size of 80 subjects per arm (240 total) provides 90% power to detect a "
        "difference of 0.6% in HbA1c between STR-4021 25 mg and placebo, assuming a standard "
        "deviation of 1.1% using a two-sided t-test at alpha=0.05. The sample size allows for "
        "up to 15% dropout. The enrollment target is 240 subjects (80 subjects per arm)."
    )

    # ─────────────────────────────────────────────────
    # SECTION 8 — SAFETY MONITORING
    # ─────────────────────────────────────────────────
    add_heading(doc, "8. Safety Monitoring", level=1)

    add_heading(doc, "8.1 Stopping Rules", level=2)
    add_paragraph(doc, "The following stopping rules apply:")
    add_bullet(doc,
        "Individual subject stopping: two consecutive HbA1c values >11.5% while on treatment "
        "will result in subject discontinuation and initiation of rescue medication."
    )
    add_bullet(doc,
        "Trial-level stopping: recommendation from the DSMB based on safety signal identified "
        "during interim safety review."
    )

    add_heading(doc, "8.2 Data Safety Monitoring Board", level=2)
    add_paragraph(doc,
        "An independent Data Safety Monitoring Board (DSMB) of three members (two clinicians "
        "and one statistician) will review unblinded safety data at pre-specified intervals. "
        "The DSMB will review at 25% and 50% of subject-years of exposure."
    )

    add_heading(doc, "8.3 Adverse Event Reporting Period", level=2)
    add_paragraph(doc,
        "The adverse event reporting period extends from first dose of study drug through 30 days "
        "after last dose (or 30 days after the follow-up visit for serious adverse events). "
        "All AEs occurring within this period will be recorded and assessed for causality."
    )

    # ─────────────────────────────────────────────────
    # SECTION 9 — ADMINISTRATIVE
    # ─────────────────────────────────────────────────
    add_heading(doc, "9. Administrative", level=1)

    add_heading(doc, "9.1 Informed Consent", level=2)
    add_paragraph(doc,
        "Informed consent will be obtained from all subjects prior to any study-related procedures. "
        "The informed consent form version to be used is ICF Version 1.0, dated 15 January 2026, "
        "consistent with Protocol Version 1.0. Any protocol amendments requiring changes to the "
        "ICF will require a new ICF version."
    )

    add_heading(doc, "9.2 Regulatory References", level=2)
    add_paragraph(doc, "This study will be conducted in accordance with the following regulatory guidelines:")
    for ref in [
        "ICH E6(R2) Good Clinical Practice",
        "ICH E9 Statistical Principles for Clinical Trials",
        "ICH E9(R1) Addendum on Estimands and Sensitivity Analysis",
        "EMA Guideline on Clinical Investigation of Medicinal Products in the Treatment of Diabetes Mellitus (CPMP/EWP/1080/00 Rev.1)",
        "21 CFR Parts 50, 54, 56, 312 (FDA)",
    ]:
        add_bullet(doc, ref)

    # NOTE: intentionally placed in an unexpected section (administrative)
    add_heading(doc, "9.3 Study Drug Note", level=2)
    add_paragraph(doc,
        "This note is placed here to test extractor robustness. The investigational drug name is "
        "STR-4021. The sponsor name is Structure Therapeutics. The protocol number STR-4021-201 "
        "should be extractable from multiple locations in this document."
    )

    # ─────────────────────────────────────────────────
    # ABBREVIATIONS
    # ─────────────────────────────────────────────────
    add_heading(doc, "Abbreviations and Key Terms", level=1)
    abbreviations = [
        "AE: Adverse Event",
        "BMI: Body Mass Index",
        "DSMB: Data Safety Monitoring Board",
        "eGFR: estimated Glomerular Filtration Rate",
        "FAS: Full Analysis Set",
        "FPG: Fasting Plasma Glucose",
        "GLP-1: Glucagon-Like Peptide-1",
        "HbA1c: Glycated Haemoglobin",
        "ICF: Informed Consent Form",
        "IRT: Interactive Response Technology",
        "MAR: Missing At Random",
        "MMRM: Mixed Model for Repeated Measures",
        "MNAR: Missing Not At Random",
        "REML: Restricted Maximum Likelihood",
        "SAE: Serious Adverse Event",
        "T2DM: Type 2 Diabetes Mellitus",
    ]
    for abbr in abbreviations:
        add_bullet(doc, abbr)

    # ─────────────────────────────────────────────────
    # SAVE
    # ─────────────────────────────────────────────────
    output_dir = Path(__file__).parent.parent / "data" / "protocols"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "synth_protocol.docx"
    doc.save(str(output_path))
    print(f"Generated: {output_path}")
    print(f"  Sections: {len(doc.paragraphs)} paragraphs")


if __name__ == "__main__":
    generate_protocol()
